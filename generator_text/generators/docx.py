import io
from docx import Document
from typing import List, Tuple, Any, IO
from .base import DataGenerator
from models.t5 import T5Model

class DocxGenerator(DataGenerator):
    """Generator for DOCX documents."""

    def __init__(self, t5_model: T5Model):
        self.t5_model = t5_model

    def generate(self, original_file: IO[bytes], num_files: int = 100, **kwargs) -> List[Tuple[str, bytes]]:
        doc = Document(original_file)
        original_paras = [p.text for p in doc.paragraphs if p.text.strip()]

        generated_files = []

        for file_idx in range(num_files):
            new_doc = Document()
            new_doc.add_heading('Synthetic Generated Document', 0)

            # Use batch paraphrasing for faster processing
            synth_paras = self.t5_model.paraphrase_batch(original_paras, batch_size=5)
            for synth_para in synth_paras:
                new_doc.add_paragraph(synth_para)

            output_buffer = io.BytesIO()
            new_doc.save(output_buffer)
            generated_files.append((f"generated_{file_idx+1}_{original_file.name}", output_buffer.getvalue()))

        return generated_files
